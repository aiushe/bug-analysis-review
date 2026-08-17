#!/usr/bin/env python3
"""Render a markdown report to a styled Word doc (python-docx).

USAGE: python3 make_report_docx.py <report.md> [out.docx]
"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if len(sys.argv) < 2:
    raise SystemExit('usage: make_report_docx.py <report.md> [out.docx]')
SRC = Path(sys.argv[1])
OUT = Path(sys.argv[2]) if len(sys.argv) > 2 else SRC.with_suffix('.docx')
PINK = RGBColor(0xD6, 0x1F, 0x69)   # accent magenta — swap for your brand colour
DARK = RGBColor(0x22, 0x22, 0x22)
HDR_BG = "D61F69"
ZEBRA = "F4E7EE"

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part; r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement('w:hyperlink'); h.set(qn('r:id'), r_id)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'),'1155CC'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u)
    r.append(rPr); t = OxmlElement('w:t'); t.text = text; r.append(t); h.append(r)
    paragraph._p.append(h)

TOKEN = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`|\[[^\]]+\]\([^)]+\))')
def add_runs(p, text):
    for part in TOKEN.split(text):
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r=p.add_run(part[2:-2]); r.bold=True
        elif part.startswith('`') and part.endswith('`'):
            r=p.add_run(part[1:-1]); r.font.name='Consolas'
        elif part.startswith('*') and part.endswith('*'):
            r=p.add_run(part[1:-1]); r.italic=True
        elif part.startswith('[') and re.match(r'\[([^\]]+)\]\(([^)]+)\)',part):
            m=re.match(r'\[([^\]]+)\]\(([^)]+)\)',part); add_hyperlink(p,m.group(2),m.group(1))
        else:
            p.add_run(part)

def main():
    lines = SRC.read_text().splitlines()
    doc = Document()
    doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(10.5)
    i=0; n=len(lines)
    while i<n:
        ln=lines[i]
        # table block
        if ln.strip().startswith('|') and i+1<n and set(lines[i+1].replace('|','').strip())<=set('-: '):
            rows=[];
            while i<n and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')]); i+=1
            header=rows[0]; body=[r for r in rows[2:]]
            t=doc.add_table(rows=1, cols=len(header)); t.style='Table Grid'
            for j,h in enumerate(header):
                cell=t.rows[0].cells[j]; cell.paragraphs[0].text=''
                r=cell.paragraphs[0].add_run(h); r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.size=Pt(9.5)
                shade(cell,HDR_BG)
            for ri,row in enumerate(body):
                cells=t.add_row().cells
                for j,val in enumerate(row[:len(header)]):
                    cells[j].paragraphs[0].text=''; add_runs(cells[j].paragraphs[0],val)
                    for para in cells[j].paragraphs:
                        for rr in para.runs: rr.font.size=Pt(9)
                    if ri%2==1: shade(cells[j],ZEBRA)
            doc.add_paragraph(); continue
        if ln.startswith('# '):
            h=doc.add_heading(level=0); r=h.add_run(ln[2:]); r.font.color.rgb=PINK
        elif ln.startswith('## '):
            h=doc.add_heading(level=1); r=h.add_run(ln[3:]); r.font.color.rgb=PINK
        elif ln.startswith('### '):
            h=doc.add_heading(level=2); r=h.add_run(ln[4:]); r.font.color.rgb=DARK
        elif ln.strip()=='---':
            doc.add_paragraph().add_run('').add_break()
        elif ln.startswith('> '):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3)
            add_runs(p,ln[2:]);
            for rr in p.runs: rr.italic=True; rr.font.color.rgb=RGBColor(0x55,0x55,0x55)
        elif ln.startswith('- '):
            p=doc.add_paragraph(style='List Bullet'); add_runs(p,ln[2:])
        elif re.match(r'^\d+\.\s',ln):
            p=doc.add_paragraph(style='List Number'); add_runs(p,re.sub(r'^\d+\.\s','',ln))
        elif ln.strip()=='':
            pass
        else:
            p=doc.add_paragraph(); add_runs(p,ln)
        i+=1
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size//1024} KB)")

if __name__=='__main__':
    main()
